from __future__ import annotations

import io
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from repopilot_guard.context import ManagedDocumentStore
from repopilot_guard.document_parser import DocumentParseError, extract_document_text


class DocumentParserTests(unittest.TestCase):
    def test_docx_is_converted_to_controlled_text_snapshot(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            source = root / "requirements.docx"
            document = Document()
            document.add_heading("订单需求", level=1)
            document.add_paragraph("查询必须按租户隔离。")
            table = document.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "接口"
            table.rows[0].cells[1].text = "权限"
            document.save(source)

            imported = ManagedDocumentStore(root / "state.sqlite").import_document(source, project_id="project-a")
            contents = imported.managed_path.read_text(encoding="utf-8")

            self.assertEqual("docx", imported.source_format)
            self.assertEqual(".txt", imported.managed_path.suffix)
            self.assertIn("订单需求", contents)
            self.assertIn("查询必须按租户隔离", contents)
            self.assertIn("接口 | 权限", contents)
            self.assertNotIn(str(source), contents)

    def test_pdf_text_is_extracted_and_encrypted_pdf_is_rejected(self) -> None:
        class Page:
            def extract_text(self) -> str:
                return "退款流程需要管理员审批"

        class Reader:
            is_encrypted = False
            pages = [Page()]

        with patch("repopilot_guard.document_parser.PdfReader", return_value=Reader()):
            self.assertEqual("退款流程需要管理员审批\n", extract_document_text(b"%PDF-test", ".pdf"))

        class EncryptedReader:
            is_encrypted = True
            pages: list[object] = []

        with patch("repopilot_guard.document_parser.PdfReader", return_value=EncryptedReader()):
            with self.assertRaisesRegex(DocumentParseError, "DOCUMENT_PDF_ENCRYPTED"):
                extract_document_text(b"%PDF-test", ".pdf")

    def test_scanned_pdf_and_invalid_docx_are_blocked(self) -> None:
        class EmptyPage:
            def extract_text(self) -> str:
                return ""

        class Reader:
            is_encrypted = False
            pages = [EmptyPage()]

        with patch("repopilot_guard.document_parser.PdfReader", return_value=Reader()):
            with self.assertRaisesRegex(DocumentParseError, "DOCUMENT_TEXT_UNAVAILABLE"):
                extract_document_text(b"%PDF-test", ".pdf")
        with self.assertRaisesRegex(DocumentParseError, "DOCUMENT_DOCX_INVALID"):
            extract_document_text(io.BytesIO(b"not-a-zip").getvalue(), ".docx")

    def test_document_secret_is_redacted_before_the_rag_snapshot(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            source = root / "integration.txt"
            source.write_text("apiKey: sk-should-not-enter-rag\n正常需求\n", encoding="utf-8")

            imported = ManagedDocumentStore(root / "state.sqlite").import_document(source, project_id="project-a")
            contents = imported.managed_path.read_text(encoding="utf-8")

            self.assertIn("[REDACTED]", contents)
            self.assertNotIn("sk-should-not-enter-rag", contents)


if __name__ == "__main__":
    unittest.main()
