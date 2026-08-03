"""受控解析用户主动导入的研发文档。"""

from __future__ import annotations

from io import BytesIO
from pathlib import PurePosixPath
from zipfile import BadZipFile, ZipFile

from docx import Document
from pypdf import PdfReader


MAX_DOCUMENT_SOURCE_BYTES = 10 * 1024 * 1024
MAX_DOCUMENT_EXTRACTED_CHARACTERS = 1_000_000
MAX_PDF_PAGES = 200
MAX_DOCX_ARCHIVE_ENTRIES = 2_048
MAX_DOCX_UNCOMPRESSED_BYTES = 32 * 1024 * 1024
SUPPORTED_DOCUMENT_EXTENSIONS = frozenset({".md", ".txt", ".pdf", ".docx"})


class DocumentParseError(ValueError):
    """以稳定错误码表示可审计的导入阻断原因。"""


def extract_document_text(raw: bytes, suffix: str) -> str:
    """将受大小限制的文件内容转换成 UTF-8 文本。"""

    normalized_suffix = suffix.lower()
    if normalized_suffix not in SUPPORTED_DOCUMENT_EXTENSIONS:
        raise DocumentParseError("UNSUPPORTED_DOCUMENT_TYPE")
    if not raw:
        raise DocumentParseError("DOCUMENT_TEXT_UNAVAILABLE")
    if len(raw) > MAX_DOCUMENT_SOURCE_BYTES:
        raise DocumentParseError("DOCUMENT_TOO_LARGE")
    if normalized_suffix in {".md", ".txt"}:
        return _plain_text(raw)
    if normalized_suffix == ".pdf":
        return _pdf_text(raw)
    return _docx_text(raw)


def _plain_text(raw: bytes) -> str:
    if b"\0" in raw:
        raise DocumentParseError("DOCUMENT_UNREADABLE")
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError as error:
        raise DocumentParseError("DOCUMENT_UNREADABLE") from error
    return _validate_extracted_text(text)


def _pdf_text(raw: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(raw), strict=False)
    except Exception as error:
        raise DocumentParseError("DOCUMENT_PDF_INVALID") from error
    if reader.is_encrypted:
        raise DocumentParseError("DOCUMENT_PDF_ENCRYPTED")
    if len(reader.pages) > MAX_PDF_PAGES:
        raise DocumentParseError("DOCUMENT_PDF_PAGE_LIMIT_EXCEEDED")
    pages: list[str] = []
    try:
        for page in reader.pages:
            pages.append(page.extract_text() or "")
    except Exception as error:
        raise DocumentParseError("DOCUMENT_PDF_TEXT_EXTRACTION_FAILED") from error
    return _validate_extracted_text("\n\n".join(pages))


def _docx_text(raw: bytes) -> str:
    _validate_docx_archive(raw)
    try:
        document = Document(BytesIO(raw))
    except Exception as error:
        raise DocumentParseError("DOCUMENT_DOCX_INVALID") from error
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _validate_extracted_text("\n".join(parts))


def _validate_docx_archive(raw: bytes) -> None:
    try:
        with ZipFile(BytesIO(raw)) as archive:
            infos = archive.infolist()
            if len(infos) > MAX_DOCX_ARCHIVE_ENTRIES:
                raise DocumentParseError("DOCUMENT_DOCX_ARCHIVE_LIMIT_EXCEEDED")
            total_size = 0
            names: set[str] = set()
            for info in infos:
                name = info.filename.replace("\\", "/")
                path = PurePosixPath(name)
                if path.is_absolute() or ".." in path.parts:
                    raise DocumentParseError("DOCUMENT_DOCX_INVALID")
                total_size += info.file_size
                names.add(name)
            if total_size > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise DocumentParseError("DOCUMENT_DOCX_ARCHIVE_LIMIT_EXCEEDED")
            if "word/document.xml" not in names:
                raise DocumentParseError("DOCUMENT_DOCX_INVALID")
    except DocumentParseError:
        raise
    except (BadZipFile, OSError) as error:
        raise DocumentParseError("DOCUMENT_DOCX_INVALID") from error


def _validate_extracted_text(text: str) -> str:
    # 不可见控制字符不能进入模型上下文。
    normalized = "".join(character for character in text if character in "\n\r\t" or ord(character) >= 32).strip()
    if not normalized:
        raise DocumentParseError("DOCUMENT_TEXT_UNAVAILABLE")
    if len(normalized) > MAX_DOCUMENT_EXTRACTED_CHARACTERS:
        raise DocumentParseError("DOCUMENT_EXTRACTED_TEXT_TOO_LARGE")
    return normalized + "\n"
